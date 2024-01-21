from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from time import sleep
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import NoSuchElementException
from openpyxl import load_workbook 
from math import sin, cos, sqrt, atan2, radians, log10
from os.path import exists
from os import chdir, mkdir
from shutil import copytree, rmtree, move
from pathlib import Path
import pyautogui as pya
from subprocess import Popen
from pyperclip import paste
from antenna import antenna
from random import randint, choice, uniform
from docx import Document,shared
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image

def distance(lat_1, lon_1, lat_2, lon_2):
    rad_earth = 6373

    lat1 = radians(lat_1)
    lon1 = radians(lon_1)
    lat2 = radians(lat_2)
    lon2 = radians(lon_2)

    a = sin(abs(lat2 - lat1) / 2)**2 + cos(lat1) * cos(lat2) * sin(abs(lon2 - lon1) / 2)**2
    c = 2 * atan2(sqrt(a), sqrt(1 - a))
    return rad_earth * c

def copy_text():
    pya.hotkey("ctrl", "c")
    sleep(0.01)
    return paste().replace(",",".")

def type(text, x, y):
    pya.doubleClick(x, y)
    pya.write(text)

def press_tab(x):
    for i in range(x):
        pya.press('tab')

def tab_write(x, text):
    for i in range(x):
        pya.press('tab')
    pya.write(text)
    
def tab_enter(x):
    for i in range(x):
        pya.press('tab')
    pya.press('enter')

def anti_tab(x):
    for i in range(x):
        pya.hotkey('shiftleft','shiftright','tab')

def set_coord(lat, lon, map_properties):
    if map_properties == 1:
        tab_enter(7)
    tab_write(8, str(lat).replace('.',','))
    tab_write(1, str(lon).replace('.',','))
    pya.press('enter')

def adjust_height(height, x, y, margem):
    for i in range(-2,3):
        pya.hotkey('backspace','backspace')
        pya.write(str(height+i))
        tab_enter(1)
        tab_enter(1)
        for j in ['10.png', '11.png', '12.png', '13.png', '14.png']:
            try:
                print(pya.locateOnScreen(j, region=[x,y-7,60,18], confidence=0.9))
                if isinstance(pya.locateOnScreen(j, region=[x,y-7,70,18], confidence=0.7), tuple):
                    return 1
            except pya.ImageNotFoundException:
                print("nao encontrado")
        for j in ['8.png','9.png']:
            try:
                print(pya.locateOnScreen(j, region=[x,y-7,60,18], confidence=0.95))
                if isinstance(pya.locateOnScreen(j, region=[x,y-7,60,18], confidence=0.8), tuple):
                    print(j)
                    return 1
            except pya.ImageNotFoundException:
                print("nao encontrado")
        anti_tab(2)
        
    return 0

def location_to_word(city, district, street, number, cep, lat, long, height):
    t = d.add_paragraph('Estado: Rio Grande do Sul')
    t = d.add_paragraph(f'Cidade: {city}')
    t = d.add_paragraph(f'Bairro: {district}')
    t = d.add_paragraph(f'Rua: {street}')
    t = d.add_paragraph(f'Número: {number}')
    t = d.add_paragraph(f'CEP: {cep}')
    t = d.add_paragraph(f'Coordenadas: {lat}, {long}')
    t = d.add_paragraph(f'Altura: {height}m\n')

def save_file():
    for i in range(8):
        pya.press('tab')
    sleep(2)
    pya.press('enter', presses=2)
    
lat = float(input("Digita tua latitude aq(usar ponto, nao virgula): "))
long = float(input("Digita tua latitude aq(usar ponto, nao virgula): "))
dist_min = float(input("Distancia minima da casa para antena: "))
tamanho_casa = int(input("Altura da sua casa ou edificio(em metros): "))

pya.useImageNotFoundException()
#lat = -29.361108173623705
#long = -50.80989409669787
#dist_min = 3.0
#tamanho_casa = 15 - 1
dist_max = 5
if tamanho_casa < 10:
    dist_max = 4.5
elif tamanho_casa < 16:
    dist_max = 8
else:
    dist_max = 5.5

#if not exists(Path.cwd() / 'chromedriver.exe'):
    #driver_service = Service(ChromeDriverManager().install())

#if exists(Path.cwd() / 'chromedriver.exe'):
    #driver_service = Service(executable_path=Path.cwd() / 'chromedriver.exe')

chrome_options = Options()
chrome_options.add_argument("--remote-debugging-port=9222")
#chrome_options.add_argument('--headless')
chrome_options.add_argument('--disable-gpu')
chrome_options.add_argument('--no-sandbox')
chrome_options.add_argument('--disable-dev-shm-usage')

# Pegar cidade
maps = webdriver.Chrome(options = chrome_options)
maps.get(f"https://www.google.com.br/maps/")
serach_bar = maps.find_element(By.NAME, "q")
serach_bar.send_keys(str(lat) + "," + str(long))
sleep(1)
serach_bar.send_keys(Keys.ENTER)
while True:
    try:
        city = maps.find_element(By.CLASS_NAME, "DkEaL")
    except NoSuchElementException:
        pass
    else:
        sleep(2)
        break
city = maps.find_element(By.CLASS_NAME, "DkEaL")
street = city.text[city.text.find(".")+2: city.text.find(',')]
number = city.text[city.text.find(',') + 2: city.text.find("-")]
city_name = city.text[city.text.find("- ")+1:]
#print(f"city_name: {city_name}")
cep = city_name[city_name.find('RS')+4 : ]
district = city_name[2 : city_name.find(",")]
city_name = city_name[city_name.find(",") + 2 : city_name.find(" -")]

#Criar uma pasta para o arquivo de radio mobile
if exists("C:\\RADIOMOBILE"):
    rmtree("C:\\RADIOMOBILE")
chdir(Path.cwd())
copytree("rmweng_func", "C:\\RADIOMOBILE")
    
Popen("C:\\RADIOMOBILE\\rmweng.exe")

sleep(1)
#internet properties
while True:
    sleep(1)
    try:
        if pya.locateOnScreen('start.png', confidence=0.9) != None:
            pass
    except pya.ImageNotFoundException:
        break

pya.click(pya.locateCenterOnScreen('new_project.png', confidence=0.72))
for i in range(4):
    pya.press('enter')

pya.click(pya.locateCenterOnScreen('internet.png', confidence=0.9))
sleep(0.7)
pya.click(pya.locateCenterOnScreen('strm.png'))
tab_write(2, "c:\RADIOMOBILE")
x, y = pya.locateCenterOnScreen('strm.png')
pya.click(x + 200, y)
try:
    while pya.locateOnScreen('best.png') == None:
        pya.write('s')
except pya.ImageNotFoundException:
    pass
for i in range(2):      
    pya.press('enter')
    
# Map properties
while True:
    try:
        print(pya.locateCenterOnScreen('map.png', confidence=0.9))
        if isinstance(pya.locateCenterOnScreen('map.png', confidence=0.9), tuple):
            break
    except pya.ImageNotFoundException:
        print("nao achou")
        pass
pya.click(pya.locateCenterOnScreen('map.png', confidence=0.9))
sleep(0.1)
tab_enter(4)
while True:
    try:
        if isinstance(pya.locateOnScreen('city.png', confidence=0.9), tuple):
            break
    except pya.ImageNotFoundException:
        pass
pya.write(city_name)
pya.press('enter')
width = choice(['1020', '820', '860', '890', '1100', '1120','900','950','970'])
tab_write(6, width)
height = choice(['860', '815', '840', '850', '830', '800','820','760','700'])
tab_write(1, height)
tab_write(14, 'strm')
tab_write(1, 'c:\RADIOMOBILE')
press_tab(6)

lat_calc = abs(float(copy_text())) - abs(lat)
press_tab(1)
long_calc = abs(float(copy_text())) - abs(long)

if lat_calc > 2 or lat_calc < -2 or long_calc < -2 or long_calc > 2:
    x,y = pya.locateCenterOnScreen("map_config.png", confidence=0.7)
    pya.click(x+200,y)
    tab_enter(4)
    set_coord(lat=lat, lon=long, map_properties=0)
    press_tab(3)

try:
    for i in pya.locateAllOnScreen('check2.png', confidence=0.8):
        pya.click(i)
except Exception:
    pass

pya.press('enter')

while True:
    try:
        if isinstance(pya.locateOnScreen('ready.png', confidence=0.9), tuple):
            break
    except pya.ImageNotFoundException:
        pass

#Network properties
sleep(4)
pya.hotkey('ctrlleft','ctrlright','n')
sleep(0.4)
pya.click(pya.locateCenterOnScreen('cont.png', confidence=0.6))
tab_enter(8)
pya.click(pya.locateCenterOnScreen('dat.png', confidence=0.6))

count = 0
try:
    for i in pya.locateAllOnScreen('select2.png', confidence=0.95):
        count += 1
        pya.click(i)
        if count >= 2:
            break
except Exception:
    pass
pya.click(pya.locateCenterOnScreen('member.png', confidence=0.7))
count = 0
try:
    for i in pya.locateAllOnScreen('check.png'):
        count += 1
except Exception:
    pass
try:
    for i in pya.locateAllOnScreen('square.png'):
        if count == 2:
            break
        pya.click(i)
        count += 1
except Exception:
    pass
pya.click(pya.locateCenterOnScreen('system.png', confidence=0.6))
tab_write(9, '1')
tab_write(1, 'yagi')
height_antenna = randint(27, 33)
tab_write(4, str(height_antenna))
pya.click(pya.locateCenterOnScreen('member.png', confidence=0.6))
pya.click(pya.locateCenterOnScreen('fixed.png', confidence=0.9))
pya.press('u')
pya.doubleClick(pya.locateCenterOnScreen('unit1.png', confidence=0.9))
pya.click(pya.locateCenterOnScreen('fixed.png', confidence=0.9))
for i in range(2):
    pya.press('u')
pya.press('enter')
tab_enter(11)
tab_write(5, '15')
tab_write(2, '10')
pya.click(pya.locateCenterOnScreen('system.png', confidence=0.6))
tab_enter(2)

# Unit properties 
pya.hotkey('ctrlleft','ctrlright','u')
set_coord(lat, long, 1)
pya.click(pya.locateCenterOnScreen('unit2.png', confidence=0.9))
pya.hotkey('enter', 'enter')

# Radio connection
pya.click(pya.locateCenterOnScreen('radio.png', confidence=0.55))
sleep(1)
tab_write(4, str(tamanho_casa))
tab_enter(2)
pya.press('enter')
pya.press('esc')

# Antenna Datasheets
file0 = load_workbook('datasheets.xlsx')
sheet_dat = file0['Plan1']

# Antenna locations
file1 = load_workbook('antenas.xlsx')
sheet = file1["Planilha1"]
dict_letters = {"A":2, "B":253,"C":525,"D":1591,"E":1657,"F":1874,"G":1979,"H":2278,"I":2302, "J":2490, "L":2543, "M":2618,
"N":2795,"O":3041,"P":3079,"Q":4770,"R":4783,"S":5017,"T":5986,"U":6292,"V":6361,"W":6592,"X":6594}
start_letter = dict_letters[city_name[:1]]

final = False
while True:
    start_letter +=1
    if start_letter == 6626:
        start_letter = 2
    if start_letter == dict_letters[city_name[:1]]:
        height_antenna = height_antenna + 1

    distance1 = distance(lat, long, float(sheet[f'G{start_letter}'].value), float(sheet[f'H{start_letter}'].value))
    if distance1 < dist_min or distance1 > dist_max:
        continue
    
    for i in sheet_dat['B1':'C64']:
        #freq_min, freq_max, gain, distance
        ant1 = antenna(int(sheet_dat['D{}'.format(i[0].row)].value), int(sheet_dat['E{}'.format(i[0].row)].value), int(sheet_dat['F{}'.format(i[0].row)].value), distance1)

        margem = ant1.margem(10, 15, i[0].value + i[1].value)
        if  margem == 0:
            continue
        print("margem:" + str(margem))
    
        
        # Network properties / systems for antennas
        pya.hotkey('ctrlleft','ctrlright','n')
        sleep(0.2)
        tab_enter(4)
        tab_write(6,str(ant1.freq_min))
        tab_write(1,str(ant1.freq_max))
        pya.click(pya.locateCenterOnScreen('system.png', confidence=0.6))
        tab_write(6, str(sheet_dat[f'B{i[0].row}'].value))
        tab_write(2, str(sheet_dat[f'C{i[0].row}'].value * -1))
        tab_write(4, str(int(ant1.gain / 2)))
        tab_write(x=2, text=str(height_antenna))
        pya.press('enter')

        # Unit properties
        pya.hotkey('ctrlleft','ctrlright','u')
        set_coord(sheet[f'G{start_letter}'].value, sheet[f'H{start_letter}'].value, 1)
        tab_enter(3)

        # Radio link
        pya.click(pya.locateCenterOnScreen('radio.png', confidence=0.6))
        sleep(1.3)
        #tab_write(14, str(ant1.freq_min))
        #sleep(0.1)
        #tab_write(1, str(ant1.freq_max))
        try:
            if isinstance(pya.locateOnScreen('(3).png', confidence=0.9), tuple):
                pya.press('esc')
                break
        except pya.ImageNotFoundException:
            pass
        try:
            if isinstance(pya.locateOnScreen('fresnel.png', confidence=0.7), tuple):
                pass
        except pya.ImageNotFoundException:
            break 
        x,y = pya.locateCenterOnScreen('fresnel.png', confidence=0.8)
        try:
            if isinstance(pya.locateOnScreen('0.png', region=[x+15,y-7,40,18], confidence=0.9), tuple):
                print('encontrado numeros worst fresnel')
                pya.press('esc')
                break
        except pya.ImageNotFoundException:
            pass
        
        located = False
        x,y = pya.locateCenterOnScreen('rx.png', confidence=0.6)
        print("-----------------------------------------")
        for j in ['10.png', '11.png', '12.png', '13.png', '14.png']:
            try:
                print(pya.locateOnScreen(j, region=[x,y-7,60,18], confidence=0.9))
                if isinstance(pya.locateOnScreen(j, region=[x,y-7,70,18], confidence=0.7), tuple):
                    located = True
                    print(j)
                    break
            except pya.ImageNotFoundException:
                print("nao encontrado")

        if located == False:
            for j in ['8.png','9.png']:
                try:
                    print(pya.locateOnScreen(j, region=[x,y-7,60,18], confidence=0.95))
                    if isinstance(pya.locateOnScreen(j, region=[x,y-7,60,18], confidence=0.8), tuple):
                        located = True
                        print(j)
                        break
                except pya.ImageNotFoundException:
                    print("nao encontrado")
        
        if located == False:
            press_tab(4)
            if adjust_height(tamanho_casa,x,y, margem) == 0:
                pya.press('esc')
                continue
        

        # Passando informacoes para word
        d = Document('base.docx')

        styles = d.styles
        style = styles.add_style('Header', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        style.font.size =shared.Pt(int(choice(['20','18','22','20','18'])))

        font_type = choice(['Arial','Times New Roman','Arial','Arial','Arial'])
        d.styles['Normal'].font.name = font_type
        d.styles['Normal'].font.size = shared.Pt(int(choice(['12','12','12','12','14'])))

        p = d.add_paragraph(choice(['Residência', 'Casa','Residência:','Residência:']))
        p.style = d.styles['Header']
        t=0

        maps.get(f"https://www.google.com.br/maps/")
        serach_bar = maps.find_element(By.NAME, "q")
        serach_bar.send_keys(str(float(sheet[f'G{start_letter}'].value)) + "," + str(float(sheet[f'H{start_letter}'].value)))
        sleep(1)
        serach_bar.send_keys(Keys.ENTER)
        while True:
            try:
                city = maps.find_element(By.CLASS_NAME, "DkEaL")
            except NoSuchElementException:
                None
            else:
                sleep(2)
                break
        city2 = maps.find_element(By.CLASS_NAME, "DkEaL")
        street2 = city2.text[city2.text.find(".")+2: city2.text.find(',')]
        number2 = city2.text[city2.text.find(',') + 2: city2.text.find("-")]
        city_name2 = city2.text[city.text.find("- "):]
        cep2 = city_name2[city_name2.find('RS')+4 : ]
        district2 = city_name2[2 : city_name2.find(",")]
        city_name2 = city_name2[city_name2.find(",") + 2 : city_name2.find(" -")]

        # Padroes Localizacoes
        pattern_city = randint(1,2)
        if pattern_city == 1:
            t = d.add_paragraph(city.text)
            t = d.add_paragraph(f'{str(lat)}, {str(long)}')
            t = d.add_paragraph(f'Altura: {tamanho_casa }m\n')
            p = d.add_paragraph(choice(['Antena', 'Torre','Torre escolhida','Torre:','Torre Escolhida:']))
            p.style = d.styles['Header']
            t = d.add_paragraph(city2.text)
            t = d.add_paragraph(str(float(sheet[f'G{start_letter}'].value)) + " " + str(float(sheet[f'H{start_letter}'].value)))
            t = d.add_paragraph(f'Altura: {height_antenna}m\n\n')

        elif pattern_city == 2:
            location_to_word(city_name, district, street, number, cep, lat, long, tamanho_casa + 1)
            p = d.add_paragraph(choice(['Antena', 'Torre','Torre escolhida','Torre:','Torre Escolhida:']))
            location_to_word(city_name2, district2, street2, number2, cep2, str(float(sheet[f'G{start_letter}'].value)), str(float(sheet[f'H{start_letter}'].value)), height_antenna)

        if randint(1,2) == 2:
            p = d.add_paragraph(f"{choice(['Distância:','Distância: '])} {str(distance1)}{choice(['Km','km'])}")
        else:
            t = d.add_paragraph(f"\nDistância: {str(distance1)}{choice(['Km','km'])}")

        # Escolha da Antena
        if pattern_city == 1:
            p = d.add_paragraph(12*'\n')
            p.style = d.styles['Header']
        elif pattern_city == 2:
            p = d.add_paragraph(7*'\n')
            p.style = d.styles['Header']

        p = d.add_paragraph(f"\n{choice(['Escolha da Antena:','Escolha da Antena','Escolhendo a Antena:', 'Escolha da Antena:','Especificações da antena:'])}")
        p.style = d.styles['Header']
        t = d.add_paragraph('\n')
        # Colocando imagem
        chdir(Path.cwd() / str(sheet_dat['G{}'.format(i[0].row)].value))
        selection = choice(['{}1.png','{}0.png','{}2.png','{}3.png','{}4.png'])
        background = Image.open('spec.jpg').convert('RGBA')
        try:
            frontImage = Image.open(selection.format('fi')).convert('RGBA')
        except FileNotFoundError:
            frontImage = None
            
        else:
            background.paste(frontImage, (randint(606, 615),randint(202, 205)), frontImage)
            background = background.crop((randint(0,100), randint(0,100), randint(760,803), randint(500,597)))
            background.save("new.png", format="png")

        if frontImage == None:
            if sheet_dat['G{}'.format(i[0].row)].value == "AM-2G15-120":
                background.crop((100, 100, 200, 150))
            background.save('new.png', format="png")
            d.add_picture('new.png', width=shared.Inches(uniform(2.5, 3.5)))
        else:
            d.add_picture("new.png", width=shared.Inches(uniform(6.3, 6.5)))

        
        t = d.add_paragraph(choice(['Escolhido o modelo {}', 'Escolhi o modelo {}','Antena Escolhida: {}','Modelo escolhido: {}', 'Escolhi a antena: {}', 'Escolhi a antena {}' ]).format(sheet_dat['G{}'.format(i[0].row)].value))
        frontImage = Image.open(selection.format('p')).convert('RGBA')
        p=d.add_paragraph(10*'\n')
        p=d.add_paragraph(choice(['Transceptor usado','Transceptor:','Escolhendo o transceptor','Escolhendo o transceptor:','Transceptor:']))
        p.style=d.styles['Header']
        background = Image.open('trans.jpg').convert('RGBA')
        if sheet_dat['G{}'.format(i[0].row)].value == "NB-2G18":
           background.paste(frontImage, (145, 69+ (17*((i[0].row-1))%16)), frontImage)
           background = background.crop((randint(0,50),randint(0,20), randint(699, 762), randint(375,432)))
        if not sheet_dat['G{}'.format(i[0].row)].value == "NB-2G18":
           background.paste(frontImage, (145, 69+ (20*((i[0].row-1))%16)), frontImage)
           background = background.crop((randint(0,10), randint(0,20), randint(700,756), randint(500,802)))
    
        background.save("new.png", format="png")

        d.add_picture('new.png', width=shared.Inches(uniform(6.5, 6.75)))

        if randint(0,6) > 0:
            t = d.add_paragraph(choice(['Escolhido: {}','Modelo: {}', 'Modelo {}', 'Modelo do transceptor: {}']).format(sheet_dat['A{}'.format(i[0].row)].value))
            if randint(0,5) > 3:
                s = choice(['{}x ','{}x: ', '{}x:'])
                t = d.add_paragraph(f"{s.format('T')}{i[0].value}")
                t = d.add_paragraph(f"{s.format('R')}-{i[1].value}")
                
        p=d.add_paragraph(6*'\n')
        p=d.add_paragraph(choice(['Cálculos:','Cálculos do Sistema','Calculos:','Cálculo da margem e fresnel:','Cálculos:']))
        spacing=choice([' {} ','{}'])
        mult=choice(['*',''])
        at=choice(['At= ','At(dB)= ','At=','At(dB)=','at(dB): ','at: '])
        t=d.add_paragraph(f"\n{at}32,5{spacing.format('+')}20{mult}log({distance1}){spacing.format('+')}20{mult}log({ant1.freq_max})")
        t=d.add_paragraph(f"{at}32,5{spacing.format('+')}{20*log10(distance1)}{spacing.format('+')}{20*log10(ant1.freq_max)}")
        t=d.add_paragraph(f"{at}{ant1.at}dB")
        t=d.add_paragraph(f"\n{choice(['Ac=','Atenuação nos cabos:','Ac(dB)=','Ac= ','Ac(dB)= '])}{ant1.ac}dB")
        gs=choice(['gs=','Gs(dBm)=','Gs=','Gs(dBm)= ','Gs= '])
        t=d.add_paragraph(f"\n{gs}{i[0].value}{spacing.format('+')}{i[1].value}")
        t=d.add_paragraph(f"{gs}{i[0].value+i[1].value}dBm")
        marg=choice(['Margem(dB)=','Margem=','Margem(dB)= ','Margem(dB)'])
        t=d.add_paragraph(f"{marg}{i[0].value+i[1].value}{spacing.format('+')}{ant1.gain}{spacing.format('-')}{ant1.ac}{spacing.format('-')}{ant1.at}")
        t=d.add_paragraph(f"{marg}{ant1.margem(10, 15, i[0].value + i[1].value)}\n\n")
        p=d.add_paragraph("{}".format(choice(['Raio de Fresnel','Raio de Fresnel:','Fresnel:','Raio de Fresnel:'])))
        fn=choice(['Rfn= ','Rfn=','Fn=','Fn= ','R'])
        lmd=choice(['λ=',' λ= ','ƛ=','ƛ= '])
        t=d.add_paragraph(f"{lmd}(3{spacing.format('*')}10^(8)){spacing.format('/')}{ant1.freq_max}{choice(['*10^(9)','Ghz','* 10^(9)'])}")
        t=d.add_paragraph(f"{lmd}{3*(10**8)/(ant1.freq_max*(10**6))}")
        distance1 = distance1 * 1000
        t=d.add_paragraph(f"{fn}√[({3*(10**8)/(ant1.freq_max*(10**6))}{spacing.format('*')}{distance1/2}{spacing.format('*')}{distance1/2})/({distance1/2}{spacing.format('+')}{distance1/2}")
        t=d.add_paragraph(f"{fn}{ant1.fresnel(distance1)}m")

        # Salvando word
        p.style = d.styles['Header']
        t.style = d.styles['Normal']
        if exists("C:\\Users\\Usuario\\Desktop\\Trabalho_radiomobile"):
            rmtree("C:\\Users\\Usuario\\Desktop\\Trabalho_radiomobile")
            mkdir("C:\\Users\\Usuario\\Desktop\\Trabalho_radiomobile")
        else:
            mkdir("C:\\Users\\Usuario\\Desktop\\Trabalho_radiomobile")
        chdir("../")
        pya.press('esc')
        pya.hotkey("ctrl", "s")
        try:
            for i in pya.locateAllOnScreen('final.png'):
                pya.click(i)
        except Exception: 
            pass
        pya.press('enter')
        pya.click(pya.locateCenterOnScreen('help.png', confidence=0.8))
        pya.press('enter')
        #sleep(2)
        #pya.click(pya.locateCenterOnScreen('file.png', confidence = 0.85))
        #sleep(0.2)
        #pya.write("C:\\Users\\Usuario\\Desktop\\Trabalho_radiomobile")
        #sleep(2)
        #tab_enter(9)
        #pya.press('enter')
        #pya.press('enter')
        chdir("C:\\Users\\Usuario\\Desktop\\Trabalho_radiomobile")
        move("C:\\RADIOMOBILE\\net1.bmp", "C:\\Users\\Usuario\\Desktop\\Trabalho_radiomobile")
        move("C:\\RADIOMOBILE\\net1.map", "C:\\Users\\Usuario\\Desktop\\Trabalho_radiomobile")
        move("C:\\RADIOMOBILE\\net1.net", "C:\\Users\\Usuario\\Desktop\\Trabalho_radiomobile")
        d.save('demo.docx')
        final = True
    if final == True:
        break
