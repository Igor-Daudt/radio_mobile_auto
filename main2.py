from openpyxl import load_workbook 
from math import sin, cos, sqrt, atan2, radians, log10
from os.path import exists
from os import chdir
from shutil import copytree, rmtree
from pathlib import Path
import pyautogui as pya
from subprocess import Popen
from pyperclip import paste
from random import randint, choice, uniform
from docx import Document,shared
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image


file0 = load_workbook('datasheets.xlsx')
sheet_dat = file0['Plan1']

# Passando informacoes para word
d = Document('base.docx')

styles = d.styles
style = styles.add_style('Header', WD_STYLE_TYPE.PARAGRAPH)
style.base_style = styles['Normal']
style.font.size =shared.Pt(20)

font_type = 'Arial'
d.styles['Normal'].font.name = font_type
d.styles['Normal'].font.size = shared.Pt(12)

p = d.add_paragraph(choice(['Residência', 'Casa','Residência:','Residência:']))
p.style = d.styles['Header']

chdir(Path.cwd() / 'AM-2G15-120')
selection = choice(['{}1.png','{}0.png','{}2.png','{}3.png','{}4.png'])
background = Image.open('spec.jpg').convert('RGBA')
       
    #background.paste(frontImage, (randint(606, 615),randint(202, 205)), frontImage)
    #background = background.crop((randint(0,100), randint(0,100), randint(760,803), randint(500,597)))
    #background.save("new.png", format="png")
#AM-2G15-120
background.crop((100, 100, 200, 150))
background.save('new.png', format="png")
d.add_picture('new.png', width=shared.Inches(uniform(2.5, 3.5)))